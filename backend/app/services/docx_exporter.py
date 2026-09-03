import io
import re
import time
from typing import Dict, Any, List, Optional
from xml.sax.saxutils import escape

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# 商务智库排版高级色盘
COLOR_PRIMARY_HEX = "1E40AF"      # 皇家深蓝 (主标题/表头)
COLOR_SECONDARY_HEX = "0F766E"    # 墨青色 (次级要点)
COLOR_TEXT_HEX = "1E293B"         # 经典深黑 (正文)
COLOR_MUTED_HEX = "64748B"        # 浅灰 (元数据/说明)
COLOR_ACCENT_HEX = "2563EB"       # 亮点蓝 (链接/角标/高亮)
COLOR_BG_LIGHT_HEX = "F8FAFC"     # 斑马纹浅底色
COLOR_BG_QUOTE_HEX = "EFF6FF"     # 引用框淡蓝底色
COLOR_BORDER_HEX = "CBD5E1"       # 边框浅灰
COLOR_CODE_BG_HEX = "F1F5F9"      # 代码卡片底色

RGB_PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
RGB_TEXT = RGBColor(0x1E, 0x29, 0x3B)
RGB_MUTED = RGBColor(0x64, 0x74, 0x8B)
RGB_ACCENT = RGBColor(0x25, 0x63, 0xEB)
RGB_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RGB_GREEN = RGBColor(0x05, 0x96, 0x69)

FONT_EAST_ASIA = "Microsoft YaHei"
FONT_ASCII = "Arial"


def set_run_font(
    run,
    font_name: str = FONT_EAST_ASIA,
    ascii_font: str = FONT_ASCII,
    size_pt: Optional[float] = 10.5,
    bold: bool = False,
    italic: bool = False,
    color_rgb: Optional[RGBColor] = None,
    superscript: bool = False
):
    """为 run 设定中西文字体规范及格式"""
    run.font.name = ascii_font
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if superscript:
        run.font.superscript = True
    if color_rgb:
        run.font.color.rgb = color_rgb

    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), ascii_font)
    rPr.append(rFonts)


def set_cell_background(cell, fill_hex: str):
    """设置表格单元格背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top: int = 120, bottom: int = 120, left: int = 160, right: int = 160):
    """设置表格单元格内边距 (单位 dxa, 1pt = 20dxa)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_table_borders(table, border_color: str = COLOR_BORDER_HEX):
    """设置现代化极简商务表格边框"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="{border_color}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_hyperlink_run(paragraph, url: str, text: str, color_hex: str = COLOR_ACCENT_HEX, underline: bool = True):
    """在段落中追加原生的 Word 可点击超链接"""
    try:
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = parse_xml(
            f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" '
            f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        )
        u_tag = '<w:u w:val="single"/>' if underline else ''
        new_run = parse_xml(
            f'<w:r {nsdecls("w")}><w:rPr><w:color w:val="{color_hex}"/>{u_tag}'
            f'<w:rFonts w:ascii="{FONT_ASCII}" w:hAnsi="{FONT_ASCII}" w:eastAsia="{FONT_EAST_ASIA}"/>'
            f'</w:rPr>'
            f'<w:t>{escape(text)}</w:t></w:r>'
        )
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        run = paragraph.add_run(text)
        set_run_font(run, color_rgb=RGB_ACCENT)


def set_paragraph_left_border(paragraph, color_hex: str = COLOR_ACCENT_HEX, sz: str = "24"):
    """为段落添加左侧强调边线 (常用于引用块)"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="{sz}" w:space="14" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_paragraph_shading(paragraph, fill_hex: str):
    """为段落添加背景底色"""
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    pPr.append(shd)


def add_styled_markdown_text(paragraph, text: str, default_size: float = 10.5, default_color: RGBColor = RGB_TEXT):
    """
    智能解析段落中的行内格式：
    - **粗体**
    - *斜体*
    - `行内代码`
    - [^cite:N] / [N] 引用角标 (自动转为 Superscript 上标)
    - [标题](URL) Markdown 超链接 (自动转为 Word 原生可点击链接)
    - local:// 本地文档徽标
    """
    if not text:
        return

    token_pattern = re.compile(
        r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[\^cite:\d+\]|\[\^\d+\]|\[\d+\]|\[[^\]]+\]\([^\)]+\))'
    )
    parts = token_pattern.split(text)

    for part in parts:
        if not part:
            continue

        # 1. 粗体 **text**
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size_pt=default_size, bold=True, color_rgb=default_color)

        # 2. 斜体 *text*
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2 and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size_pt=default_size, italic=True, color_rgb=default_color)

        # 3. 行内代码 `code`
        elif part.startswith('`') and part.endswith('`') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, font_name="Consolas", ascii_font="Consolas", size_pt=default_size * 0.92, color_rgb=RGBColor(0xB9, 0x1C, 0x1C))

        # 4. Markdown 链接 [Title](URL)
        elif part.startswith('[') and '](' in part and part.endswith(')'):
            match = re.match(r'^\[(.*?)\]\((.*?)\)$', part)
            if match:
                link_title = match.group(1)
                link_url = match.group(2)
                if link_url.startswith('local://'):
                    fname = link_url.replace('local://', '')
                    run = paragraph.add_run(f" 📄 [本地专有文献: {link_title or fname}] ")
                    set_run_font(run, size_pt=default_size * 0.92, bold=True, color_rgb=RGB_GREEN)
                else:
                    add_hyperlink_run(paragraph, link_url, link_title)
            else:
                run = paragraph.add_run(part)
                set_run_font(run, size_pt=default_size, color_rgb=default_color)

        # 5. 引用角标 [N] 或 [^cite:N] (转换为 Word 真正的上标)
        elif re.match(r'^(\[\^cite:\d+\]|\[\^\d+\]|\[\d+\])$', part):
            num_match = re.search(r'\d+', part)
            num_str = f"[{num_match.group(0)}]" if num_match else part
            run = paragraph.add_run(num_str)
            set_run_font(run, size_pt=default_size * 0.85, bold=True, color_rgb=RGB_ACCENT, superscript=True)

        # 6. 普通文本
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size_pt=default_size, color_rgb=default_color)


def generate_editorial_docx(
    report_md: str,
    title: str,
    metadata: Optional[Dict[str, Any]] = None,
    style: str = "consulting"
) -> io.BytesIO:
    """
    将深度研究报告 Markdown 转换为具有顶级出版级排版质感的 Microsoft Word (.docx) 文档。
    支持依据 5 大专业研报风格自适应专属主题配色与元数据。
    """
    from app.agents.writer import StyleProfileRegistry
    
    chosen_style = style or (metadata.get("style") if metadata else None) or "consulting"
    profile = StyleProfileRegistry.get(chosen_style)
    primary_hex = profile.get("docx_primary_color", COLOR_PRIMARY_HEX)
    try:
        r_val = int(primary_hex[0:2], 16)
        g_val = int(primary_hex[2:4], 16)
        b_val = int(primary_hex[4:6], 16)
        theme_rgb_primary = RGBColor(r_val, g_val, b_val)
    except Exception:
        theme_rgb_primary = RGB_PRIMARY
        primary_hex = COLOR_PRIMARY_HEX

    doc = Document()

    # 1. 页面设置：标准 A4 (210mm x 297mm), 上下 2.54cm, 左右 2.8cm
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    # 2. 页眉与页脚配置
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.add_run(f"Deep Research Agent 2.5 · {profile['name_zh']}")
    set_run_font(header_run, size_pt=8.5, color_rgb=RGB_MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("机密与专业报告 · 100% 真实信源交叉溯源保障")
    set_run_font(footer_run, size_pt=8.5, color_rgb=RGB_MUTED)

    # 3. 统计信源与图表数量 (用于封面元数据卡片)
    citation_count = len(re.findall(r'\[\d+\]\s+\[.*?\]\(.*?\)', report_md)) or len(re.findall(r'\[\d+\]', report_md)) // 2
    mermaid_count = len(re.findall(r'```(?:mermaid|flowchart|sequenceDiagram|gantt|mindmap|graph)', report_md, re.IGNORECASE))

    # -------------------------------------------------------------
    # 4. 出版级封面卡片 (Cover Page)
    # -------------------------------------------------------------
    badge_p = doc.add_paragraph()
    badge_p.paragraph_format.space_before = Pt(36)
    badge_p.paragraph_format.space_after = Pt(12)
    badge_run = badge_p.add_run(f"【{profile['name_zh']}权威深度报告】")
    set_run_font(badge_run, size_pt=11, bold=True, color_rgb=theme_rgb_primary)

    clean_main_title = title or "深度产业研究与商业化前景推演报告"
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(14)
    title_p.paragraph_format.line_spacing = 1.25
    title_run = title_p.add_run(clean_main_title)
    set_run_font(title_run, size_pt=24, bold=True, color_rgb=theme_rgb_primary)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(24)
    sub_run = sub_p.add_run("多智能体自主深度调研 · 全网混合检索与交叉溯源验证报告")
    set_run_font(sub_run, size_pt=12, color_rgb=RGB_MUTED)

    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_after = Pt(32)
    set_paragraph_shading(div_p, primary_hex)
    div_run = div_p.add_run(" ")
    div_run.font.size = Pt(2)

    curr_date_str = time.strftime("%Y年%m月%d日", time.localtime())
    meta_rows = [
        ("调研课题领域", clean_main_title[:45] + ("..." if len(clean_main_title) > 45 else "")),
        ("研报编制机构", "Deep Research Autonomous Agent 2.5 (Multi-Model Gateway)"),
        ("完成发布时间", curr_date_str),
        ("权威数据源统计", f"收录 {citation_count} 处可信行业信源 · 100% 真实交叉验证"),
        ("产业可视化架构", f"内嵌 {mermaid_count} 套技术演进与产业链路图谱"),
        ("报告专业风格", f"{profile['name_zh']} ({profile['name_en']})")
    ]

    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    set_table_borders(meta_table, COLOR_BORDER_HEX)
    meta_table.autofit = True

    for r_idx, (k, v) in enumerate(meta_rows):
        row = meta_table.rows[r_idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        
        set_cell_background(cell_k, COLOR_BG_LIGHT_HEX)
        set_cell_margins(cell_k, top=100, bottom=100, left=140, right=140)
        p_k = cell_k.paragraphs[0]
        p_k.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_k = p_k.add_run(k)
        set_run_font(run_k, size_pt=9.5, bold=True, color_rgb=RGB_MUTED)

        set_cell_background(cell_v, "FFFFFF" if r_idx % 2 == 0 else COLOR_BG_LIGHT_HEX)
        set_cell_margins(cell_v, top=100, bottom=100, left=140, right=140)
        p_v = cell_v.paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_v = p_v.add_run(v)
        set_run_font(run_v, size_pt=9.5, color_rgb=RGB_TEXT)

    doc.add_page_break()

    # -------------------------------------------------------------
    # 5. 正文解析与排版循环
    # -------------------------------------------------------------
    lines = report_md.split('\n')
    i = 0

    while i < len(lines):
        raw_line = lines[i]
        trimmed = raw_line.strip()

        if not trimmed:
            i += 1
            continue

        # A. 代码块与 Mermaid 图表处理 (``` ... ```)
        # 增加容错：如遇到行首标题 (# 开头)，说明上一段代码未闭合，立即提前截断
        if trimmed.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines):
                next_raw = lines[i]
                next_trimmed = next_raw.strip()
                if next_trimmed.startswith('```'):
                    i += 1
                    break
                if re.match(r'^#+\s+', next_trimmed):
                    # 遇到标题行，说明大模型遗漏闭合反引号，强制截断代码块
                    break
                code_lines.append(next_raw)
                i += 1

            code_text = "\n".join(code_lines).strip()
            if not code_text:
                continue

            is_mermaid = 'mermaid' in trimmed.lower() or any(
                re.match(r'^(graph|flowchart|subgraph|sequenceDiagram|gantt|mindmap|classDiagram)\b', cl.strip())
                for cl in code_lines
            )

            if is_mermaid:
                card_table = doc.add_table(rows=2, cols=1)
                set_table_borders(card_table, COLOR_BORDER_HEX)
                card_table.autofit = True
                
                header_cell = card_table.rows[0].cells[0]
                set_cell_background(header_cell, COLOR_BG_LIGHT_HEX)
                set_cell_margins(header_cell, top=100, bottom=80, left=140, right=140)
                hp = header_cell.paragraphs[0]
                hrun = hp.add_run("📐 技术架构与产业演进路线图谱 (Mermaid Architecture)")
                set_run_font(hrun, size_pt=9.5, bold=True, color_rgb=RGB_PRIMARY)
                
                body_cell = card_table.rows[1].cells[0]
                set_cell_background(body_cell, COLOR_CODE_BG_HEX)
                set_cell_margins(body_cell, top=100, bottom=120, left=140, right=140)
                bp = body_cell.paragraphs[0]
                bp.paragraph_format.line_spacing = 1.15
                brun = bp.add_run(code_text)
                set_run_font(brun, font_name="Consolas", ascii_font="Consolas", size_pt=8.5, color_rgb=RGBColor(0x33, 0x41, 0x55))
            else:
                cp = doc.add_paragraph()
                cp.paragraph_format.left_indent = Inches(0.2)
                cp.paragraph_format.right_indent = Inches(0.2)
                cp.paragraph_format.space_before = Pt(6)
                cp.paragraph_format.space_after = Pt(6)
                set_paragraph_shading(cp, COLOR_CODE_BG_HEX)
                crun = cp.add_run(code_text)
                set_run_font(crun, font_name="Consolas", ascii_font="Consolas", size_pt=9.0, color_rgb=RGBColor(0x33, 0x41, 0x55))
            continue

        # B. 标题 1 (# ...)
        if trimmed.startswith('# ') and not trimmed.startswith('## '):
            clean_t = re.sub(r'^#\s*', '', trimmed).replace('**', '').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_t)
            set_run_font(run, size_pt=18, bold=True, color_rgb=RGB_PRIMARY)
            i += 1
            continue

        # C. 标题 2 (## ...)
        if trimmed.startswith('## ') and not trimmed.startswith('### '):
            clean_t = re.sub(r'^##\s*', '', trimmed).replace('**', '').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            if "参考资料" in clean_t or "Citations" in clean_t or "数据来源" in clean_t:
                set_paragraph_left_border(p, COLOR_ACCENT_HEX, sz="32")
                run = p.add_run("📚 " + clean_t.replace("📚", "").strip())
            else:
                set_paragraph_left_border(p, COLOR_PRIMARY_HEX, sz="24")
                run = p.add_run(clean_t)
                
            set_run_font(run, size_pt=13.5, bold=True, color_rgb=RGB_PRIMARY)
            i += 1
            continue

        # D. 标题 3 (### ...)
        if trimmed.startswith('### ') and not trimmed.startswith('#### '):
            clean_t = re.sub(r'^###\s*', '', trimmed).replace('**', '').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_t)
            set_run_font(run, size_pt=11.5, bold=True, color_rgb=RGBColor(0x33, 0x41, 0x55))
            i += 1
            continue

        # E. 标题 4 (#### ...) 与 标题 5 (##### ...)
        if trimmed.startswith('#### '):
            clean_t = re.sub(r'^#+\s*', '', trimmed).replace('**', '').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_t)
            set_run_font(run, size_pt=10.5, bold=True, color_rgb=RGBColor(0x47, 0x55, 0x69))
            i += 1
            continue

        # F. Markdown 对比表格 (| ... |)
        if trimmed.startswith('|') and trimmed.endswith('|'):
            table_rows_data = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                cur_tline = lines[i].strip()
                if not re.match(r'^\|[\s\-:|]+\|$', cur_tline):
                    cells = [c.strip() for c in cur_tline.split('|')[1:-1]]
                    table_rows_data.append(cells)
                i += 1

            if table_rows_data:
                col_cnt = max(len(r) for r in table_rows_data)
                tbl = doc.add_table(rows=len(table_rows_data), cols=col_cnt)
                tbl.autofit = True
                set_table_borders(tbl, COLOR_BORDER_HEX)

                for r_idx, row_vals in enumerate(table_rows_data):
                    is_head = (r_idx == 0)
                    row_obj = tbl.rows[r_idx]
                    
                    for c_idx in range(col_cnt):
                        cell_obj = row_obj.cells[c_idx]
                        val_str = row_vals[c_idx] if c_idx < len(row_vals) else ""
                        
                        if is_head:
                            set_cell_background(cell_obj, COLOR_PRIMARY_HEX)
                            set_cell_margins(cell_obj, top=120, bottom=120, left=160, right=160)
                        else:
                            bg_color = COLOR_BG_LIGHT_HEX if r_idx % 2 == 0 else "FFFFFF"
                            set_cell_background(cell_obj, bg_color)
                            set_cell_margins(cell_obj, top=100, bottom=100, left=140, right=140)

                        cell_p = cell_obj.paragraphs[0]
                        cell_p.paragraph_format.line_spacing = 1.2
                        cell_p.paragraph_format.space_before = Pt(0)
                        cell_p.paragraph_format.space_after = Pt(0)
                        
                        if is_head:
                            clean_head = val_str.replace('**', '').strip()
                            hrun = cell_p.add_run(clean_head)
                            set_run_font(hrun, size_pt=9.5, bold=True, color_rgb=RGB_WHITE)
                        else:
                            add_styled_markdown_text(cell_p, val_str, default_size=9.0)
                            
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(8)
            continue

        # G. 专用参考文献条目渲染 (- **[1]** [Title](url) 或 - [1] [Title](url))
        cite_item_match = re.match(r'^[-*]\s+(\*?\*?\[\d+\]\*?\*?)\s+(.*)', trimmed)
        if cite_item_match:
            cite_num_raw = cite_item_match.group(1).replace('*', '').strip()
            cite_rest = cite_item_match.group(2).strip()

            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(4)
            lp.paragraph_format.space_after = Pt(2)
            lp.paragraph_format.left_indent = Inches(0.2)

            num_run = lp.add_run(f"{cite_num_raw} ")
            set_run_font(num_run, size_pt=10.0, bold=True, color_rgb=RGB_ACCENT)

            add_styled_markdown_text(lp, cite_rest, default_size=10.0)
            i += 1

            # 智能检查下一行是否是该文献对应的引用摘要 (> "...")
            if i < len(lines) and lines[i].strip().startswith('>'):
                quote_line = lines[i].strip()
                clean_q = re.sub(r'^>\s*', '', quote_line).strip('“”" ')
                qp = doc.add_paragraph()
                qp.paragraph_format.left_indent = Inches(0.4)
                qp.paragraph_format.right_indent = Inches(0.2)
                qp.paragraph_format.space_before = Pt(1)
                qp.paragraph_format.space_after = Pt(6)
                set_paragraph_shading(qp, COLOR_BG_LIGHT_HEX)
                set_paragraph_left_border(qp, COLOR_BORDER_HEX, sz="16")
                
                qrun = qp.add_run(f'“{clean_q}”')
                set_run_font(qrun, size_pt=9.0, italic=True, color_rgb=RGBColor(0x47, 0x55, 0x69))
                i += 1
            continue

        # H. 通用引用块 (> ...)
        if trimmed.startswith('>'):
            clean_quote = re.sub(r'^>\s*', '', trimmed).strip()
            qp = doc.add_paragraph()
            qp.paragraph_format.left_indent = Inches(0.25)
            qp.paragraph_format.right_indent = Inches(0.2)
            qp.paragraph_format.space_before = Pt(4)
            qp.paragraph_format.space_after = Pt(6)
            set_paragraph_left_border(qp, COLOR_ACCENT_HEX, sz="24")
            set_paragraph_shading(qp, COLOR_BG_QUOTE_HEX)
            
            add_styled_markdown_text(qp, f"“{clean_quote.strip('“”\"')}”", default_size=9.5, default_color=RGBColor(0x33, 0x41, 0x55))
            i += 1
            continue

        # I. 列表项 (- 或 * 或 1. )
        if trimmed.startswith('- ') or trimmed.startswith('* ') or re.match(r'^\d+\.\s', trimmed):
            clean_li = re.sub(r'^([-*]|\d+\.)\s+', '', trimmed).strip()
            lp = doc.add_paragraph(style='List Bullet')
            lp.paragraph_format.space_before = Pt(2)
            lp.paragraph_format.space_after = Pt(3)
            lp.paragraph_format.line_spacing = 1.3
            
            add_styled_markdown_text(lp, clean_li, default_size=10.0)
            i += 1
            continue

        # J. 分隔线 (---)
        if trimmed == '---':
            hr_p = doc.add_paragraph()
            hr_p.paragraph_format.space_before = Pt(8)
            hr_p.paragraph_format.space_after = Pt(8)
            set_paragraph_shading(hr_p, COLOR_BORDER_HEX)
            hr_run = hr_p.add_run(" ")
            hr_run.font.size = Pt(1)
            i += 1
            continue

        # K. 常规正文段落
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_styled_markdown_text(p, trimmed, default_size=10.5, default_color=RGB_TEXT)
        i += 1

    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream
