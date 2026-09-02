import io
import re
import urllib.parse
from typing import List, Dict, Any, Optional
from fastapi import APIRouter, HTTPException, status, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sse_starlette.sse import EventSourceResponse

from app.services.task_manager import task_manager, TaskStatus
from app.agents.state import ChapterOutline
from app.core.config import call_llm
from app.tools.doc_parser import parse_uploaded_document

router = APIRouter()

class CreateTaskRequest(BaseModel):
    query: str = Field(..., description="调研课题 / 核心命题", min_length=2)
    depth: str = Field("standard", description="调研深度 (quick | standard | deep)")
    style: str = Field("consulting", description="报告风格 (consulting | academic | executive)")
    auto_approve_outline: bool = Field(True, description="是否自动批准大纲")
    max_iterations: int = Field(2, description="最大反思循环轮数", ge=1, le=5)
    local_documents: Optional[List[Dict[str, Any]]] = Field(None, description="上传的本地私有文档库切片")

class CreateTaskResponse(BaseModel):
    task_id: str
    status: str
    message: str

class ChapterOutlineItem(BaseModel):
    chapter_num: int
    title: str
    focus: str
    search_queries: List[str] = []
    extracted_facts: List[str] = []

class ApproveOutlineRequest(BaseModel):
    outline: List[ChapterOutlineItem] = Field(..., description="确认或修改后的章节大纲列表")

class ChatRequest(BaseModel):
    question: str = Field(..., description="用户追问内容")
    report_context: Optional[str] = Field(None, description="研报正文与事实上下文")
    task_id: Optional[str] = Field(None, description="所属任务ID")

class ChatResponse(BaseModel):
    answer: str

class ExportDocxRequest(BaseModel):
    report: str = Field(..., description="研报 Markdown 文本")
    title: str = Field("深度行业研究报告", description="研报标题")

@router.post("/upload")
async def upload_local_document(file: UploadFile = File(...)):
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="上传文件为空")
    
    parsed = parse_uploaded_document(file.filename or "uploaded_file", file_bytes)
    return parsed

@router.post("/tasks", response_model=CreateTaskResponse, status_code=status.HTTP_201_CREATED)
async def create_research_task(req: CreateTaskRequest):
    task_id = task_manager.create_task(
        user_query=req.query,
        research_depth=req.depth,
        report_style=req.style,
        auto_approve_outline=req.auto_approve_outline,
        max_iterations=req.max_iterations,
        local_documents=req.local_documents
    )
    return {
        "task_id": task_id,
        "status": TaskStatus.PENDING,
        "message": "调研任务已创建，可通过 /stream 接口订阅实时事件流"
    }

@router.get("/tasks/{task_id}")
async def get_task_detail(task_id: str):
    task = task_manager.get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail=f"任务 {task_id} 不存在")
    return task

@router.get("/tasks")
async def list_all_tasks():
    return task_manager.list_tasks()

@router.get("/tasks/{task_id}/stream")
async def stream_task_events(task_id: str):
    task = task_manager.get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail=f"任务 {task_id} 不存在")
        
    return EventSourceResponse(
        task_manager.subscribe_stream(task_id),
        media_type="text/event-stream"
    )

@router.post("/tasks/{task_id}/approve_outline")
async def approve_task_outline(task_id: str, req: ApproveOutlineRequest):
    raw_outline = [item.model_dump() for item in req.outline]
    success = task_manager.approve_outline(task_id, raw_outline)
    if not success:
        task = task_manager.get_task(task_id)
        if not task:
            raise HTTPException(status_code=404, detail="任务不存在")
        raise HTTPException(
            status_code=400,
            detail=f"任务当前状态为 {task['status']}，并非等待大纲确认状态"
        )
    return {"message": "大纲已确认，Agent 已恢复执行后续全网调研！"}

@router.post("/tasks/{task_id}/cancel")
async def cancel_research_task(task_id: str):
    success = task_manager.cancel_task(task_id)
    if not success:
        raise HTTPException(status_code=404, detail=f"任务 {task_id} 不存在")
    return {"message": f"任务 {task_id} 已取消"}

@router.post("/chat", response_model=ChatResponse)
async def chat_with_report(req: ChatRequest):
    report_text = req.report_context or ""
    if not report_text and req.task_id:
        task = task_manager.get_task(req.task_id)
        if task and "state" in task:
            report_text = task["state"].get("final_report", "")

    if not report_text:
        report_text = "（暂无具体研报上下文，请基于常识与事实回答）"

    prompt = f"""
你是一位资深的首席研究专家与智库分析师。
以下是已完成的深度研究报告全文及事实库：

【研报全文与事实来源】：
{report_text[:8000]}

【用户的具体追问 / 划词深挖请求】：
{req.question}

请严格基于上述报告中的论据、数据指标与核心逻辑，为用户提供条理清晰、有深度、具有建设性的专业解答。请直接回答，条理清晰，使用 Markdown 格式。
"""
    try:
        answer = call_llm(
            prompt,
            system_prompt="你是一位客观、严谨、具有深厚洞察力的智库研究分析师。",
            temperature=0.3
        )
        return {"answer": answer}
    except Exception as e:
        return {"answer": f"调用大模型追问失败: {str(e)}"}

def add_markdown_paragraph_to_doc(doc, text: str, is_bullet: bool = False):
    """
    将包含 **粗体**、*斜体* 的 markdown 文本转换为 Word 原生带格式段落
    """
    p = doc.add_paragraph(style='List Bullet' if is_bullet else 'Normal')
    
    # 解析粗体与斜体
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endsWith('**') if hasattr(part, 'endsWith') else part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)

@router.post("/export/docx")
async def export_report_to_docx(req: ExportDocxRequest):
    """
    【出版级 Word 文档导出】
    完整解析各级标题 (# 到 #####)、表格、加粗、斜体与引用，彻底剥离原始 Markdown 符号
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 文档大标题
        title_p = doc.add_heading(req.title, level=0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        lines = req.report.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            trimmed = line.strip()
            
            if not trimmed:
                i += 1
                continue
                
            # 1. 标题 1 (# ...)
            if trimmed.startswith('# ') and not trimmed.startswith('## '):
                clean_t = re.sub(r'#+\s*', '', trimmed).replace('**', '').strip()
                doc.add_heading(clean_t, level=1)
                i += 1
            # 2. 标题 2 (## ...)
            elif trimmed.startswith('## ') and not trimmed.startswith('### '):
                clean_t = re.sub(r'#+\s*', '', trimmed).replace('**', '').strip()
                doc.add_heading(clean_t, level=2)
                i += 1
            # 3. 标题 3 (### ...)
            elif trimmed.startswith('### ') and not trimmed.startswith('#### '):
                clean_t = re.sub(r'#+\s*', '', trimmed).replace('**', '').strip()
                doc.add_heading(clean_t, level=3)
                i += 1
            # 4. 标题 4 (#### ...)
            elif trimmed.startswith('#### ') and not trimmed.startswith('##### '):
                clean_t = re.sub(r'#+\s*', '', trimmed).replace('**', '').strip()
                doc.add_heading(clean_t, level=4)
                i += 1
            # 5. 标题 5 (##### ...)
            elif trimmed.startswith('##### '):
                clean_t = re.sub(r'#+\s*', '', trimmed).replace('**', '').strip()
                doc.add_heading(clean_t, level=5)
                i += 1
            # 6. Markdown 表格 (| ... |)
            elif trimmed.startswith('|') and trimmed.endswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                    t_line = lines[i].strip()
                    # 跳过对齐行 |--|--|
                    if not re.match(r'^\|[\s\-:|]+\|$', t_line):
                        cells = [c.strip().replace('**', '') for c in t_line.split('|')[1:-1]]
                        table_lines.append(cells)
                    i += 1
                
                if table_lines:
                    col_count = max(len(r) for r in table_lines)
                    table = doc.add_table(rows=len(table_lines), cols=col_count)
                    table.style = 'Table Grid'
                    for r_idx, row_cells in enumerate(table_lines):
                        for c_idx, cell_value in enumerate(row_cells):
                            if c_idx < col_count:
                                cell = table.cell(r_idx, c_idx)
                                cell.text = cell_value
                                if r_idx == 0:
                                    # 表头加粗
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            run.bold = True
                continue
            # 7. 列表项 (- 或 * )
            elif trimmed.startswith('- ') or trimmed.startswith('* '):
                clean_item = trimmed[2:].strip()
                add_markdown_paragraph_to_doc(doc, clean_item, is_bullet=True)
                i += 1
            # 8. 引用块 (> ...)
            elif trimmed.startswith('>'):
                clean_quote = re.sub(r'^>\s*', '', trimmed).strip()
                p = doc.add_paragraph(clean_quote)
                p.paragraph_format.left_indent = Inches(0.25)
                for run in p.runs:
                    run.italic = True
                i += 1
            # 9. 分隔线 (---)
            elif trimmed == '---':
                i += 1
            # 10. 常规段落
            else:
                add_markdown_paragraph_to_doc(doc, trimmed, is_bullet=False)
                i += 1
                
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        filename = f"{req.title}.docx"
        encoded_filename = urllib.parse.quote(filename)
        
        return StreamingResponse(
            docx_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出 Word 失败: {str(e)}")
