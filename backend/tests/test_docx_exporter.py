import io
import pytest
from docx import Document
from app.services.docx_exporter import generate_editorial_docx

def test_generate_editorial_docx_complete_structure():
    sample_md = """# 全球固态电池深度研报

## 执行摘要与核心战略洞察

固态电池正处于实验室向中试跃迁的关键期[^cite:1]。以下为量化对比：

| 路线维度 | 硫化物路线 | 氧化物路线 |
| :--- | :--- | :--- |
| 室温离子电导率 | 10⁻³ S/cm | 10⁻⁴ S/cm |
| 电化学稳定窗口 | 0.4~2.5V | 0~5.0V |

```mermaid
graph TD
    A[液态电芯] --> B[半固态]
    B --> C[全固态电芯]
```

### 关键工艺与工程壁垒
- 等静压工艺需要在300MPa超高压下成型[2]
- 界面接触阻抗高需通过复合涂层解决

> "全固态电池商业化不仅依赖材料突破，更依赖良率与规模化设备国产化。"

---

## 📚 参考资料与可信数据来源 (Citations & Sources)

- **[1]** [中国汽车技术研究中心 2025 固态电池白皮书](https://example.com/paper.pdf)
  > "全固态电芯在中试阶段面临硫化锂提纯与干法涂布良率双重考验。"
- **[2]** [本地专属企业调研内部专报](local://battery_report.docx)
  > "公司中试线计划于2026年完成车规级电芯装车测试。"
"""

    stream = generate_editorial_docx(sample_md, "全球固态电池深度研报")
    assert isinstance(stream, io.BytesIO)
    data = stream.getvalue()
    assert len(data) > 10000

    # 读取并验证生成的 docx 结构
    doc = Document(io.BytesIO(data))
    
    # 1. 验证封面与正文分页
    assert len(doc.sections) >= 1
    sec = doc.sections[0]
    assert sec.page_width.inches == pytest.approx(8.27, 0.05)
    assert sec.page_height.inches == pytest.approx(11.69, 0.05)

    # 2. 验证元数据表格和数据表格
    # Table 0: 封面元数据卡片
    # Table 1: 路线维度对比表格
    # Table 2: Mermaid 架构图卡片
    assert len(doc.tables) >= 3

    meta_table = doc.tables[0]
    assert len(meta_table.rows) == 6
    assert meta_table.rows[0].cells[0].text == "调研课题领域"

    # 3. 验证表格表头与数据
    data_table = doc.tables[1]
    assert data_table.rows[0].cells[0].text.strip() == "路线维度"
    assert data_table.rows[1].cells[1].text.strip() == "10⁻³ S/cm"

    # 4. 验证 Mermaid 架构图卡片已独立存在且无遗漏
    mermaid_table = doc.tables[2]
    assert "Mermaid Architecture" in mermaid_table.rows[0].cells[0].text
    assert "graph TD" in mermaid_table.rows[1].cells[0].text

    # 5. 验证段落与超链接存在
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "执行摘要与核心战略洞察" in full_text
    assert "全固态电池商业化不仅依赖材料突破" in full_text
    assert "本地专有文献" in full_text
